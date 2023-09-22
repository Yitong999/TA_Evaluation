from Survey import Survey
from TA import TA
from TAs import TAs

import os
import numpy as np
import pandas as pd
import openpyxl
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

STANDARD = {
    'Strongly Disagree': 1,
    'Disagree': 2,
    'Neutral': 3,
    'Agree': 4,
    'Strongly Agree': 5
}


def read_surveys():
    # TODO: iterate through all files in **Work_Space** folder

    NUM_Q = 5
    df = pd.read_excel('Work_Space/row_data.xlsx', header=None)
    num_rows = df.shape[0]

    evaluation_start_idx = find_evaluation_start_idx(df)
    evaluation_end_idx = evaluation_start_idx + NUM_Q - 1
    print('evaluation_end_col: ', evaluation_end_idx)
    #
    #
    #
    # Iterate the loop to read the cell values
    surveys = []
    for row in range(1, num_rows, 1):
        row_data = df.iloc[row].tolist()

        name = find_TA_name(row_data, evaluation_start_idx)
        scores = find_scores(row_data, evaluation_start_idx, evaluation_end_idx)
        comment = find_comment(row_data, evaluation_end_idx + 1)

        if name != 'empty' and len(scores) == NUM_Q:
            survey = Survey(name, scores, comment)
            surveys.append(survey)

    return surveys


def eval_surveys(surveys):
    TAs_result = TAs()

    for survey in surveys:
        TAs_result.add_survey(survey)

    TAs_result.update_all()
    return TAs_result


def find_TA_name(row_data, max):
    for i in range(1, max, 1):
        item = row_data[i]
        if isinstance(item, str):
            return item

    return 'empty'


def find_scores(row_data, start, end):
    scores_row = row_data[start: end + 1]

    scores = [STANDARD.get(evaluation) for evaluation in scores_row if evaluation in STANDARD]

    return scores


def find_comment(row_data, index):
    return row_data[index]


def find_evaluation_start_idx(dataframe):
    for col_number, col_title in enumerate(dataframe.iloc[0]):
        if col_title == 'Q1 well prepared':
            return col_number

    return -1

def write_to_docx(TA_name, data, comments):
    # Create a new Document
    doc = Document()

    image_path = 'src/header.jpg'
    font_name = 'Times New Roman'

    # doc.sections[0].header.add_picture(image_path, Inches(0), Inches(0), width=Inches(6))
    header = doc.sections[0].header
    par = header.paragraphs[0]

    logo_run = par.add_run()
    logo_run.add_picture(image_path, width=Inches(6.5))
    # header.add_picture(image_path, width=Inches(6))

    doc_header = doc.add_heading('Summary of TA evaluations for ' + TA_name)
    doc_header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for run in doc_header.runs:
        run.font.size = Pt(18)
        run.font.name = font_name
        run.font.color.rgb = RGBColor(0, 0, 0)

    title_1 = doc.add_paragraph()
    title_1.add_run('\nSummary of Student responses')

    for run in title_1.runs:
        run.font.size = Pt(12)
        run.font.name = font_name
        run.bold = True

    # Determine the number of rows and columns in the 2D array
    num_rows = len(data)
    num_cols = len(data[0])

    # Add a table with the same number of rows and columns as the 2D array
    table = doc.add_table(rows=num_rows, cols=num_cols)


    table.style = 'TableGrid'
    # Set the width of the table columns
    table.autofit = True

    # for i in range(num_cols):
    #     table.columns[i].width = Pt(100)  # You can adjust the column width as needed

    # Set the alignment of text within the cells
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Populate the table with data from the 2D array
    for i in range(num_rows):
        row = table.rows[i]
        for j in range(num_cols):
            cell = row.cells[j]
            cell.text = str(data[i][j])  # Convert to string if data is not already a string

    # Merge cells
    table.cell(0, 0).merge(table.cell(1, 1))
    table.cell(0, 2).merge(table.cell(0, 6))
    table.cell(0, 7).merge(table.cell(1, 7))
    table.cell(0, 8).merge(table.cell(1, 8))

    description = doc.add_paragraph()
    description.add_run(
        '* 5 point Likert scale: 1 = Strongly disagree (SD), 2 = Disagree (D), 3 = Neutral (N), 4 = Agree (A), 5 = Strongly agree (SA)')
    description.add_run('\n')
    description.add_run('\n')

    for run in description.runs:
        run.font.name = font_name
        run.font.size = Pt(10)

    title_2 = doc.add_paragraph()
    title_2.add_run('Student comments')


    for run in title_2.runs:
        run.font.size = Pt(12)
        run.font.name = font_name
        run.bold = True


    paragraph = doc.add_paragraph()
    for index, each in enumerate(comments, start=1):
        label = f'{index}. '
        paragraph.add_run(label).bold = True
        paragraph.add_run(str(each) + '\n')

    for run in paragraph.runs:
        run.font.name = font_name

    # Save the document to a file
    doc.save(f'Output/evaluation_{TA_name}.docx')

    print('Table document created successfully.')


if __name__ == '__main__':
    surveys = read_surveys()
    eval_results = eval_surveys(surveys)
    TA_sets = eval_results.TAs_set

    for TA in TA_sets:
        num_rows = 7
        num_cols = 9

        # Initialize a 7x9 array with empty strings
        empty_array = [['' for _ in range(num_cols)] for _ in range(num_rows)]

        empty_array[2][0] = 'Q1'
        empty_array[3][0] = 'Q2'
        empty_array[4][0] = 'Q3'
        empty_array[5][0] = 'Q4'
        empty_array[6][0] = 'Q5'

        empty_array[2][1] = 'The TA was well prepared:'
        empty_array[3][1] = 'The TA was helpful:'
        empty_array[4][1] = 'The TA was considerate of students:'
        empty_array[5][1] = 'The TA was easily understood:'
        empty_array[6][1] = 'The TA was an effective instructor:'

        empty_array[0][2] = 'Number of Responses'

        empty_array[1][2] = 'SD'
        empty_array[1][3] = 'D'
        empty_array[1][4] = 'N'
        empty_array[1][5] = 'A'
        empty_array[1][6] = 'SA'

        empty_array[1][7] = 'Total'
        empty_array[1][8] = 'Average*'

        for i in range(5):
            for j in range(5):
                empty_array[2 + i][2 + j] = str(round(TA.scores_table[i][j]))

        for i in range(5):
            empty_array[2 + i][7] = str(round(TA.total[i]))

        for i in range(5):
            value = TA.avg[i]
            empty_array[2 + i][8] = str(f'{value:.1f}')

        write_to_docx(TA.name, empty_array, TA.comments)

